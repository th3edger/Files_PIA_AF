#CREAR 100 ARCHIVOS WORD#
import random
from docx import Document
from docx.shared import Cm

nombres = (
    "RecetaPaciente_","RecetaLab._","Nota_",
    "HistoriaClinica_","ResumenDeCaso_","Pedidos_",
    "ListaMedica_","Medicamentos_","Antibioticos_",
    "FormatoHist.Clinica_","ListaEstServicioSoc","Estudiantes_",
    "ListEnfermeras_","Nomina_","Presupuestos_",
    "ListadoHerramientas_","ListaInstrumentos_","PedidosMeds_",
    "PedidosMateriales_","Materiales_"
)

def elegir_materia():
    materia = random.choice(nombres)
    materia.encode()
    return materia

def main():
    f = open('chanchuyo.txt', 'w')
    contador = 1
    while contador < 26:
        # Creación del documento
        document = Document()

        core_properties = document.core_properties
        core_properties.author = 'Edgar Alberto Pérez Arredondo'

        # Añadimos un titulo al documento, a nivel 0
        titulo_documento = "Documento"+str(contador)+" creado con Python"
        document.add_heading(titulo_documento, 0)

        # Añadimos un párrafo
        p = document.add_paragraph('El contenido de los párrafos se añadir en varias líneas. ')
        p.add_run('Pudiéndose configurar que el texto tenga formato tipo ')
        p.add_run('negrita').bold = True
        p.add_run(' o ')
        p.add_run('itálica.').italic = True

        # Para indicar subtitulo se indica el nivel 1
        document.add_heading('Subtitulo', level=1)

        document.add_paragraph('Ahora se puede crear una enumeración')
        document.add_paragraph('Uno', style='List Number')
        document.add_paragraph('Dos', style='List Number')
        document.add_paragraph('Tres', style='List Number')

        document.add_paragraph('O viñetas')
        document.add_paragraph('Manzana', style='List Bullet') 
        document.add_paragraph('Pera', style='List Bullet')
        document.add_paragraph('Naranja', style='List Bullet')
        document.add_paragraph('Piña', style='List Bullet') 
        document.add_paragraph('Melón', style='List Bullet')
        document.add_paragraph('Sandía', style='List Bullet')

        # Tablas
        document.add_heading('Tablas', level=1)

        data = (('Manzana', 12), ('Pera', 5), ('Naranja', 12), ('Piña', 89), ('Melón', 65), ('Sandía', 7))

        table = document.add_table(rows=1, cols=2)

        table.rows[0].cells[0].text = 'Fruta'
        table.rows[0].cells[1].text = 'Cantidad'

        for prod, numbr in data:
            row_cells = table.add_row().cells
            row_cells[0].text = prod
            row_cells[1].text = str(numbr)

        nombre_documento = elegir_materia()
        nombre_documento = nombre_documento + str(contador)+".doc"

        f.write(f'un Documento licito es: {nombre_documento}\n')

        document.save(nombre_documento)

        contador += 1

    f.close


if __name__ == '__main__':
	main()
print("Hecho!!!")